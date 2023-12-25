from docx2pdf import convert
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from num2words import num2words

document = Document()

# Add a section to the document
section = document.sections[0]

# Add a header to the section
header = section.header

# Add a paragraph to the header
paragraph_header = header.paragraphs[0]
run = paragraph_header.add_run()

# Add the image to the header
image_path = 'templates/header_SF.png'  # Replace with the path to your image file
SF_logo = run.add_picture(image_path)  # Adjust width and height as needed
paragraph_header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


# Add a footer to the section
footer = section.footer

# Add a paragraph to the footer
paragraph_footer = footer.paragraphs[0]
run = paragraph_footer.add_run()
paragraph_footer.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
paragraph_footer.paragraph_format.first_line_indent = Inches(-0.8)

# Add the image to the footer and resize it
image_path = 'templates/footer_SF.png'  # Replace with the path to your footer image file
run.add_picture(image_path, width=Inches(8), height=Inches(0.8))  # Adjust width and height as needed

# Offer letter Variables

offereddate = '24-12-2023'
candidatename = 'Naveen Kumar'
candidateAddress = "Madanapalli"
worklocation = "Bangalore,KA"
position = "Territory Head (Project Coordinator)"
salary = 563247
reporting = "Genral Manager"
joining = '30-12-2023'


words = str(num2words(salary, lang='en_IN')).capitalize()

page1 =f'''


{offereddate}

{candidatename}
{candidateAddress}

Dear {candidatename},

We are pleased to make you an offer for the position of {position} in our organization. 

You shall be paid a compensation of Rs.{salary}/- ({words} Only) per annum. The details are set out in the annexure.

The detailed appointment letter along with the Employment Contract will be given to you on your joining the company.

You will be based at {worklocation} and you will report for your work and assignment to the 
{reporting}.

You are requested to sign the duplicate of this letter as a token of your acceptance and report for joining on or before {joining}

We welcome you to our organization and wish you a very successful career with us.'''

p = document.add_paragraph(str(page1))

document.add_paragraph(f"""
Yours Sincerely,
for Sambhav Foundation 						Accepted""")
image_path = 'templates/gayatri_shankar.png'  # Replace with the path to your image file
document.add_picture(image_path)
document.add_paragraph(f"""Gayathri Shanker				                   	  	{candidatename}
Head - People & Culture""")
document.add_page_break()
document.add_paragraph()
table1 = document.add_table(rows=1, cols=1)
table1.cell(0, 0).text = "Annexure"
table2 = document.add_table(rows=11, cols=2)

# Add text to specific cells in the table

table2.cell(0, 0).text = "Name"
table2.cell(1, 0).text = "Designation"
table2.cell(2, 0).text = "Location"
table2.cell(3, 0).text = "Components"
table2.cell(4, 0).text = "Basic Salary"
table2.cell(5, 0).text = "House Rent Allowance"
table2.cell(6, 0).text = "Special Allowance"
table2.cell(7, 0).text = "Gross Salary"
table2.cell(8, 0).text = "Employer Provident Fund"
table2.cell(9, 0).text = "Employer Insurance (GPA+GMC)"
table2.cell(10, 0).text = "Total CTC"

table2.cell(0, 1).text = candidatename
table2.cell(1, 1).text = position
table2.cell(2, 1).text = worklocation
table2.cell(3, 1).text = "CTC Per Annum (in INR)"
table2.cell(4, 1).text = '100'
table2.cell(5, 1).text = '10'
table2.cell(6, 1).text = '20'
table2.cell(7, 1).text = '30'
table2.cell(8, 1).text = '2'
table2.cell(9, 1).text = '10'
table2.cell(10, 1).text = "Total CTC"
# Table Formating
table1.style = 'Table Grid'
table2.style = 'Table Grid'

# paragraph = document.add_paragraph()
# paragraph.style = 'ListBullet'
# paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# paragraph.paragraph_format.left_indent = Pt(36)

points = ["""The entitlements are subject to company policies / procedures / guidelines that may be issued / modified from time to time. All perquisites and benefits including reimbursements are subject to Income Tax provisions, which may be applicable, including taxation on perquisite value.""",
          """These entitlements shall cease upon the termination of your employment with Sambhav Foundation. These entitlements may also cease if you need to take long-termed personal leave of absence. You will need to check with your local HR team for details.""",
          """Insurance Coverage:
    •	Medical Group Insurance - Rs. 5.0 Lakh per annum (Employee, spouse and 2 dependent children upto 21 years of age will be covered under floater policy)
    •	Personal Accidental Insurance - Rs. 7.0 Lakh per annum coverage""",
          """Tax is applicable as per norms""",
          """Your compensation remains strictly private and confidential between the company and you, this must not be disclosed to others."""]

document.add_paragraph()
for i in points:
    paragraph = document.add_paragraph(i)
    paragraph.style = 'ListNumber'

document.add_paragraph(f"""
Yours Sincerely,
for Sambhav Foundation 						Accepted""")
image_path = 'templates/gayatri_shankar.png'  # Replace with the path to your image file
document.add_picture(image_path)
document.add_paragraph(f"""Gayathri Shanker				                   	  	{candidatename}
Head - People & Culture""")

# Save the document
document.save(str(candidatename)+".docx")
convert(str(candidatename)+".docx")